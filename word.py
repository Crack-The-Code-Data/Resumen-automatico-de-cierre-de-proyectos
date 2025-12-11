from io import BytesIO
from collections import OrderedDict
from typing import List, Optional, Tuple, Dict, Any
import pandas as pd

# Librerías para manejo de documentos Word (python-docx)
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL


def crear_documento_a4() -> Document:
    """
    Crea un documento Word con tamaño de página A4 y márgenes de 1 pulgada.
    """
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    return doc


def guardar_documento(doc: Document, ruta_salida: str) -> None:
    doc.save(ruta_salida)


def agregar_titulo(doc: Document, texto: str, nivel: int) -> None:
    # Paleta de colores corporativos sobrios
    COLOR_TITULO = RGBColor(0x2E, 0x3F, 0x5F)  # Azul marino oscuro
    COLOR_SUBTITULO = RGBColor(0x4F, 0x4F, 0x4F)  # Gris oscuro

    if nivel == 1:
        # Título principal - Nivel 1
        titulo = doc.add_heading(level=1)
        run = titulo.add_run(texto.upper())
        run.font.name = 'Lora'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_TITULO
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        titulo.paragraph_format.space_before = Pt(18)
        titulo.paragraph_format.space_after = Pt(12)

        # Agregar línea decorativa inferior
        p = titulo._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2E3F5F')
        pBdr.append(bottom)

    elif nivel == 2:
        # Subtítulo importante - Nivel 2
        titulo = doc.add_heading(level=2)
        run = titulo.add_run(texto)
        run.font.name = 'Lora'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_SUBTITULO
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        titulo.paragraph_format.space_before = Pt(14)
        titulo.paragraph_format.space_after = Pt(8)

        # Subrayado decorativo
        p = titulo._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'D3D3D3')
        pBdr.append(bottom)

    elif nivel == 3:
        # Subtítulo secundario - Nivel 3
        titulo = doc.add_heading(level=3)
        run = titulo.add_run(texto)
        run.font.name = 'Lora'
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_SUBTITULO
        run.font.italic = True
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        titulo.paragraph_format.space_before = Pt(10)
        titulo.paragraph_format.space_after = Pt(4)

    else:
        # Para niveles inferiores
        parrafo = doc.add_paragraph(texto)
        parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        run = parrafo.runs[0]
        run.font.name = 'Segoe UI Light'
        run.font.size = Pt(8)
        run.font.underline = True
        run.font.bold = True


def agregar_parrafo(doc: Document, texto: str) -> None:
    parrafo = doc.add_paragraph(texto)
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = parrafo.runs[0]
    run.font.name = 'Segoe UI Light'
    run.font.size = Pt(8)


def insertar_figura(doc: Document, figura, titulo: Optional[str] = None, pie: Optional[str] = None, ancho_cm: Optional[float] = None, alto_cm: Optional[float] = None) -> None:
    if titulo:
        agregar_titulo(doc, titulo, 3)
    imagen_stream = BytesIO()
    figura.savefig(imagen_stream, format='png', bbox_inches='tight')
    imagen_stream.seek(0)

    p = doc.add_paragraph()
    run = p.add_run()

    # Determinar dimensiones de la imagen
    if ancho_cm is not None and alto_cm is not None:
        run.add_picture(imagen_stream, width=Cm(ancho_cm), height=Cm(alto_cm))
    elif ancho_cm is not None:
        run.add_picture(imagen_stream, width=Cm(ancho_cm))
    elif alto_cm is not None:
        run.add_picture(imagen_stream, height=Cm(alto_cm))
    else:
        run.add_picture(imagen_stream, width=Inches(5.5))

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    imagen_stream.close()

    if pie:
        pie_p = doc.add_paragraph(pie)
        pie_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = pie_p.runs[0]
        run.font.name = 'Segoe UI Light'
        run.font.size = Pt(6)
        run.font.bold = True
        run.font.italic = True


def set_cell_width(cell, width_inches: float) -> None:
    """
    Establece el ancho de una celda en pulgadas.
    """
    width_twips = int(width_inches * 1440)
    cell.width = Inches(width_inches)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Eliminar cualquier w:tcW anterior
    for child in tcPr.findall(qn('w:tcW')):
        tcPr.remove(child)

    # Crear nuevo elemento de ancho
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def insertar_tabla(doc: Document, df, titulo: Optional[str] = None):
    if titulo:
        agregar_titulo(doc, titulo, 3)

    tabla = doc.add_table(rows=1, cols=len(df.columns))
    tabla.style = 'Table Grid'
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    ancho_total = 6.0
    ancho_columna = ancho_total / len(df.columns)

    # Encabezados
    hdr_cells = tabla.rows[0].cells
    for i, col_name in enumerate(df.columns):
        cell = hdr_cells[i]
        cell.text = str(col_name)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(6.5)
        run.font.name = 'Segoe UI Light'
        set_cell_width(cell, ancho_columna)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Filas de datos
    for _, row in df.iterrows():
        row_cells = tabla.add_row().cells
        for i, value in enumerate(row):
            cell = row_cells[i]
            cell.text = str(value)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(7)
            set_cell_width(cell, ancho_columna)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return tabla


def insertar_tabla_con_merge(doc: Document, df, titulo: Optional[str] = None, group_cols: Optional[List[str]] = None):
    if titulo:
        agregar_titulo(doc, titulo, 3)

    tabla = doc.add_table(rows=1, cols=len(df.columns))
    tabla.style = 'Table Grid'
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    ancho_total = 6.0
    ancho_columna = ancho_total / len(df.columns)

    # Encabezados
    hdr_cells = tabla.rows[0].cells
    for i, col in enumerate(df.columns):
        cell = hdr_cells[i]
        cell.text = str(col)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(6.5)
        run.font.name = 'Segoe UI Light'
        set_cell_width(cell, ancho_columna)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Filas de datos
    for _, row in df.iterrows():
        row_cells = tabla.add_row().cells
        for i, val in enumerate(row):
            cell = row_cells[i]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(7)
            set_cell_width(cell, ancho_columna)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if group_cols:
        col2idx = {col: idx for idx, col in enumerate(df.columns)}
        sizes = OrderedDict()
        prev_key = None
        for key_vals in df[group_cols].itertuples(index=False, name=None):
            if key_vals == prev_key:
                sizes[key_vals] += 1
            else:
                sizes[key_vals] = 1
                prev_key = key_vals

        current_row = 1
        for key_vals, size in sizes.items():
            if size > 1:
                for col in group_cols:
                    c_idx = col2idx[col]
                    start = tabla.cell(current_row, c_idx)
                    end = tabla.cell(current_row + size - 1, c_idx)
                    for r in range(current_row + 1, current_row + size):
                        tabla.cell(r, c_idx).text = ''
                    start.merge(end)
                    start.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    start.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            current_row += size

    return tabla


def insertar_salto_pagina(doc: Document) -> None:
    doc.add_page_break()


def agregar_viñetas(
    doc: Document,
    items: List[str],
    nivel: int = 1,
    espacio_antes: Pt = Pt(4),
    espacio_despues: Pt = Pt(4),
) -> None:
    """
    Inserta una lista usando guiones '-' como viñetas.
    """
    indent_por_nivel = Pt(12)

    for texto in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = espacio_antes
        p.paragraph_format.space_after = espacio_despues
        p.paragraph_format.left_indent = indent_por_nivel * (nivel - 1)

        run = p.add_run(f"- {texto}")
        run.font.name = 'Segoe UI Light'
        run.font.size = Pt(8)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def insertar_en_posicion(
    doc: Document,
    funcion_contenido,
    *args,
    posicion: str = 'final',
    **kwargs,
) -> None:
    """
    Inserta contenido generado por una función en una posición específica del documento.
    posicion: 'inicio', 'final' o 'index:<n>'
    """
    doc_temp = Document()
    funcion_contenido(doc_temp, *args, **kwargs)

    elementos_temp = list(doc_temp.element.body)
    body = doc.element.body

    if posicion == 'inicio':
        for elem in reversed(elementos_temp):
            body.insert(0, elem)
    elif posicion == 'final':
        for elem in elementos_temp:
            body.append(elem)
    elif posicion.startswith('index:'):
        idx = int(posicion.split(':')[1])
        for i, elem in enumerate(elementos_temp):
            body.insert(idx + i, elem)
    else:
        raise ValueError("La posición debe ser 'inicio', 'final' o 'index:<n>'")


def insertar_indice(doc: Document, titulo: str = "Índice") -> None:
    agregar_titulo(doc, titulo, 1)
    p = doc.add_paragraph()
    run = p.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_after = Pt(6)


def agregar_advertencia_actualizacion(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run("⚠️ Al abrir este documento, recuerde actualizar los campos (índice, referencias cruzadas, etc.).")
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(12)


def mostrar_contenido(doc: Document) -> None:
    print("Índice | Tipo   | Contenido resumido")
    print("--------------------------------------")

    idx_parrafo = 0
    idx_tabla = 0

    for i, elem in enumerate(doc.element.body):
        tag = elem.tag.split('}')[-1]

        if tag == 'p':
            parrafo = doc.paragraphs[idx_parrafo]
            texto = parrafo.text.strip().replace('\n', ' ')
            print(f"{i:<6} | Párrafo | '{texto[:60]}'")
            idx_parrafo += 1

        elif tag == 'tbl':
            print(f"{i:<6} | Tabla   | [Tabla con {len(doc.tables[idx_tabla].rows)} filas]")
            idx_tabla += 1

        else:
            print(f"{i:<6} | Otro    | Etiqueta: {tag}")


def mostrar_contenido_posicional(doc: Document, buscar: Optional[str] = None) -> List[int]:
    idx_parrafo = 0
    posiciones_encontradas: List[int] = []

    for i, elem in enumerate(doc.element.body):
        tag = elem.tag.split('}')[-1]

        if tag == 'p':
            parrafo = doc.paragraphs[idx_parrafo]
            texto = parrafo.text.strip().replace('\n', ' ')
            if buscar and buscar.lower() in texto.lower():
                posiciones_encontradas.append(i)
            idx_parrafo += 1

    return posiciones_encontradas


def reemplazar_parrafo(original: Paragraph, nuevo: Paragraph) -> None:
    original._element.getparent().replace(original._element, nuevo._element)


def numerar_titulos_existentes(doc: Document) -> None:
    contador = {1: 0, 2: 0, 3: 0}
    reemplazos: List[Tuple[Paragraph, Paragraph]] = []

    for _, parrafo in enumerate(doc.paragraphs):
        estilo = parrafo.style.name.strip()
        if estilo.startswith("Heading"):
            try:
                nivel = int(estilo.split()[-1])
            except (ValueError, IndexError):
                continue

            if nivel in contador:
                contador[nivel] += 1
                for deeper in range(nivel + 1, 4):
                    contador[deeper] = 0

                if nivel == 1:
                    numeracion = f"{contador[1]}."
                elif nivel == 2:
                    numeracion = f"{contador[1]}.{contador[2]}"
                else:
                    numeracion = f"{contador[1]}.{contador[2]}.{contador[3]}"

                texto = parrafo.text.strip()
                if not texto.startswith(numeracion):
                    doc_temp = Document()
                    agregar_titulo(doc_temp, f"{numeracion} {texto}", nivel)
                    nuevo_parrafo = doc_temp.paragraphs[0]
                    reemplazos.append((parrafo, nuevo_parrafo))

    for original, nuevo in reemplazos:
        reemplazar_parrafo(original, nuevo)


# ============================================================================
# FORMATO CONDICIONAL PARA TABLAS - Funciones auxiliares
# ============================================================================


def set_cell_background(cell, color: str):
    """
    Establece color de fondo de una celda.
    
    Args:
        cell: Celda de la tabla
        color: Color en formato hexadecimal sin '#' (ej: 'FF0000' para rojo)
    """
    # Normalizar color y aplicar sombreado de celda usando propiedades correctas
    hex_color = (color or '').replace('#', '').upper()
    if len(hex_color) != 6:
        return

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Eliminar sombreado previo si existe
    existing_shd = tcPr.findall(qn('w:shd'))
    for shd in existing_shd:
        tcPr.remove(shd)

    # Crear nuevo sombreado con atributos necesarios
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


# ==================== FORMATOS CONDICIONALES ====================

def formato_filas_alternadas(tabla, color1='FFFFFF', color2='F0F0F0', desde_fila=1):
    """
    Aplica colores alternados a las filas (zebra striping).
    
    Args:
        tabla: Tabla de Word retornada por insertar_tabla()
        color1: Color para filas pares (hexadecimal sin #)
        color2: Color para filas impares (hexadecimal sin #)
        desde_fila: Desde qué fila empezar (1=después del encabezado)
    
    Ejemplo:
        >>> tabla = insertar_tabla(doc, df, "Mi Tabla")
        >>> formato_filas_alternadas(tabla, 'FFFFFF', 'E8F4F8')
    """
    for i, row in enumerate(tabla.rows[desde_fila:], start=desde_fila):
        color = color1 if (i - desde_fila) % 2 == 0 else color2
        for cell in row.cells:
            set_cell_background(cell, color)


def formato_valores_positivos_negativos(tabla, df, columnas=None, 
                                       color_positivo='C6EFCE', 
                                       color_negativo='FFC7CE',
                                       color_cero='FFFFFF'):
    """
    Colorea celdas según si el valor es positivo, negativo o cero.
    
    Args:
        tabla: Tabla de Word retornada por insertar_tabla()
        df: DataFrame original usado para crear la tabla
        columnas: Lista de nombres de columnas a aplicar formato (None=todas)
        color_positivo: Color para valores > 0
        color_negativo: Color para valores < 0
        color_cero: Color para valores = 0
    
    Ejemplo:
        >>> tabla = insertar_tabla(doc, df, "Ventas")
        >>> formato_valores_positivos_negativos(tabla, df, columnas=['Variación', 'Margen'])
    """
    if columnas is None:
        columnas = df.columns
    
    col_indices = [list(df.columns).index(col) for col in columnas if col in df.columns]
    
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        for col_idx in col_indices:
            try:
                valor = float(row.iloc[col_idx])
                if valor > 0:
                    color = color_positivo
                elif valor < 0:
                    color = color_negativo
                else:
                    color = color_cero
                set_cell_background(tabla.rows[i].cells[col_idx], color)
            except (ValueError, TypeError):
                pass


def formato_por_umbral(tabla, df, columna, umbrales, colores):
    """
    Colorea celdas según umbrales definidos.
    
    Args:
        tabla: Tabla de Word
        df: DataFrame original
        columna: Nombre de la columna a aplicar formato
        umbrales: Lista de valores umbral (ordenados de menor a mayor)
        colores: Lista de colores (len(umbrales) + 1)
    
    Ejemplo:
        >>> # Rojo<50, Amarillo 50-80, Verde claro 80-90, Verde oscuro >90
        >>> tabla = insertar_tabla(doc, df, "Rendimiento")
        >>> formato_por_umbral(tabla, df, 'Porcentaje', 
        ...                    umbrales=[50, 80, 90],
        ...                    colores=['FFC7CE', 'FFEB9C', 'C6EFCE', '92D050'])
    """
    if columna not in df.columns:
        return
    
    col_idx = list(df.columns).index(columna)
    
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        try:
            valor = float(row.iloc[col_idx])
            color_idx = 0
            for umbral in umbrales:
                if valor >= umbral:
                    color_idx += 1
            set_cell_background(tabla.rows[i].cells[col_idx], colores[color_idx])
        except (ValueError, TypeError):
            pass


def formato_top_bottom(tabla, df, columna, top_n=3, color_top='C6EFCE', 
                       bottom_n=3, color_bottom='FFC7CE'):
    """
    Resalta los N valores más altos y más bajos de una columna.
    
    Args:
        tabla: Tabla de Word
        df: DataFrame original
        columna: Nombre de la columna
        top_n: Cantidad de valores máximos a resaltar
        color_top: Color para valores máximos
        bottom_n: Cantidad de valores mínimos a resaltar
        color_bottom: Color para valores mínimos
    
    Ejemplo:
        >>> tabla = insertar_tabla(doc, df, "Ranking")
        >>> formato_top_bottom(tabla, df, 'Ventas', top_n=5, bottom_n=3)
    """
    if columna not in df.columns:
        return
    
    col_idx = list(df.columns).index(columna)
    
    try:
        valores_numericos = pd.to_numeric(df[columna], errors='coerce')
        top_indices = valores_numericos.nlargest(top_n).index
        bottom_indices = valores_numericos.nsmallest(bottom_n).index
        
        for i, idx in enumerate(df.index, start=1):
            if idx in top_indices:
                set_cell_background(tabla.rows[i].cells[col_idx], color_top)
            elif idx in bottom_indices:
                set_cell_background(tabla.rows[i].cells[col_idx], color_bottom)
    except:
        pass


def formato_escala_color(tabla, df, columna, Vmin=None, Vmax=None, color_min='FFC7CE', 
                        color_max='C6EFCE', color_medio='FFEB9C'):
    """
    Aplica gradiente de color según el valor (mapa de calor).
    
    Args:
        tabla: Tabla de Word
        df: DataFrame original
        columna: Nombre de la columna
        color_min: Color para valor mínimo (hex)
        color_max: Color para valor máximo (hex)
        color_medio: Color para valor medio (hex)
    
    Ejemplo:
        >>> tabla = insertar_tabla(doc, df, "Temperatura")
        >>> formato_escala_color(tabla, df, 'Temperatura', 
        ...                      color_min='5B9BD5',  # Azul
        ...                      color_max='FF0000',  # Rojo
        ...                      color_medio='FFFF00') # Amarillo
    """
    if columna not in df.columns:
        return
    
    col_idx = list(df.columns).index(columna)
    
    try:
        valores = pd.to_numeric(df[columna], errors='coerce')

        if Vmin:
            min_val = Vmin
        else:
            min_val = valores.min()
        if Vmax:
            max_val = Vmax
        else:
            max_val = valores.max()
        
        def hex_to_rgb(hex_color):
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        
        def rgb_to_hex(rgb):
            return '%02x%02x%02x' % rgb
        
        rgb_min = hex_to_rgb(color_min)
        rgb_medio = hex_to_rgb(color_medio)
        rgb_max = hex_to_rgb(color_max)
        
        for i, (_, row) in enumerate(df.iterrows(), start=1):
            try:
                valor = float(row.iloc[col_idx])
                if pd.notna(valor):
                    # Normalizar valor entre 0 y 1
                    if max_val != min_val:
                        ratio = (valor - min_val) / (max_val - min_val)
                    else:
                        ratio = 0.5
                    
                    # Interpolación de color
                    if ratio < 0.5:
                        r = int(rgb_min[0] + (rgb_medio[0] - rgb_min[0]) * (ratio * 2))
                        g = int(rgb_min[1] + (rgb_medio[1] - rgb_min[1]) * (ratio * 2))
                        b = int(rgb_min[2] + (rgb_medio[2] - rgb_min[2]) * (ratio * 2))
                    else:
                        r = int(rgb_medio[0] + (rgb_max[0] - rgb_medio[0]) * ((ratio - 0.5) * 2))
                        g = int(rgb_medio[1] + (rgb_max[1] - rgb_medio[1]) * ((ratio - 0.5) * 2))
                        b = int(rgb_medio[2] + (rgb_max[2] - rgb_medio[2]) * ((ratio - 0.5) * 2))
                    
                    color = rgb_to_hex((r, g, b))
                    set_cell_background(tabla.rows[i].cells[col_idx], color)
            except (ValueError, TypeError):
                pass
    except:
        pass


def formato_columnas_especificas(tabla, df, columnas_colores):
    """
    Colorea columnas completas (excepto encabezado).
    
    Args:
        tabla: Tabla de Word
        df: DataFrame original
        columnas_colores: Dict {nombre_columna: color_hex}
    
    Ejemplo:
        >>> tabla = insertar_tabla(doc, df, "Datos")
        >>> formato_columnas_especificas(tabla, df, {
        ...     'Ventas': 'E7E6E6',
        ...     'Margen': 'D9E1F2',
        ...     'Beneficio': 'E2EFDA'
        ... })
    """
    for columna, color in columnas_colores.items():
        if columna in df.columns:
            col_idx = list(df.columns).index(columna)
            for i in range(1, len(tabla.rows)):
                set_cell_background(tabla.rows[i].cells[col_idx], color)


def formato_contiene_texto(tabla, df, columna, texto_color):
    """
    Colorea celdas que contienen cierto texto.
    
    Args:
        tabla: Tabla de Word
        df: DataFrame original
        columna: Nombre de la columna
        texto_color: Dict {texto_buscar: color_hex}
    
    Ejemplo:
        >>> tabla = insertar_tabla(doc, df, "Estados")
        >>> formato_contiene_texto(tabla, df, 'Estado', {
        ...     'Aprobado': 'C6EFCE',
        ...     'Rechazado': 'FFC7CE',
        ...     'Pendiente': 'FFEB9C',
        ...     'En Revisión': 'BDD7EE'
        ... })
    """
    if columna not in df.columns:
        return
    
    col_idx = list(df.columns).index(columna)
    
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        valor = str(row.iloc[col_idx])
        for texto, color in texto_color.items():
            if texto.lower() in valor.lower():
                set_cell_background(tabla.rows[i].cells[col_idx], color)
                break


def formato_resaltar_duplicados(tabla, df, columna, color='FFEB9C'):
    """
    Resalta valores duplicados en una columna.
    
    Args:
        tabla: Tabla de Word
        df: DataFrame original
        columna: Nombre de la columna
        color: Color para duplicados
    
    Ejemplo:
        >>> tabla = insertar_tabla(doc, df, "IDs")
        >>> formato_resaltar_duplicados(tabla, df, 'ID', color='FFC7CE')
    """
    if columna not in df.columns:
        return
    
    col_idx = list(df.columns).index(columna)
    valores_duplicados = df[df.duplicated(subset=[columna], keep=False)][columna].unique()
    
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        if row.iloc[col_idx] in valores_duplicados:
            set_cell_background(tabla.rows[i].cells[col_idx], color)


def formato_encabezado_personalizado(tabla, color_fondo='2E3F5F', color_texto='FFFFFF'):
    """
    Aplica formato personalizado al encabezado de la tabla.
    
    Args:
        tabla: Tabla de Word
        color_fondo: Color de fondo del encabezado (hex)
        color_texto: Color del texto (hex) - formato 'RRGGBB'
    
    Ejemplo:
        >>> tabla = insertar_tabla(doc, df)
        >>> formato_encabezado_personalizado(tabla, '2E3F5F', 'FFFFFF')
    """
    for cell in tabla.rows[0].cells:
        set_cell_background(cell, color_fondo)
        # Cambiar color del texto
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                r, g, b = int(color_texto[0:2], 16), int(color_texto[2:4], 16), int(color_texto[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)


# ==================== PALETA DE COLORES CORPORATIVOS ====================

class PaletaColores:
    """Paleta de colores predefinida para formato condicional."""
    
    # Semáforo
    VERDE = 'C6EFCE'
    VERDE_OSCURO = '92D050'
    AMARILLO = 'FFEB9C'
    NARANJA = 'FFD966'
    ROJO = 'FFC7CE'
    ROJO_OSCURO = 'FF6B6B'
    
    # Azules corporativos
    AZUL_CLARO = 'D9E1F2'
    AZUL_MEDIO = 'B4C7E7'
    AZUL_OSCURO = '8EA9DB'
    AZUL_MARINO = '2E3F5F'
    
    # Grises
    GRIS_MUY_CLARO = 'F8F9FA'
    GRIS_CLARO = 'F0F0F0'
    GRIS_MEDIO = 'D0D0D0'
    GRIS_OSCURO = '6C757D'
    
    # Otros
    MORADO_CLARO = 'E2EFDA'
    ROSA = 'FCE4D6'
    CYAN = 'D0F0FD'
    BLANCO = 'FFFFFF'

# ============================================================================
# CLASE DOCUMENTBUILDER - Patrón Builder para construcción fluida de documentos
# ============================================================================

class DocumentBuilder:
    """
    Constructor fluido de documentos Word con configuración centralizada.

    Permite crear documentos mediante encadenamiento de métodos y mantiene
    un historial de operaciones para debugging.

    Ejemplo básico:
        >>> builder = DocumentBuilder()
        >>> builder.titulo("Mi Reporte", 1) \\
        ...        .parrafo("Introducción al documento.") \\
        ...        .vinetas(["Punto 1", "Punto 2"]) \\
        ...        .guardar("reporte.docx")

    Ejemplo con configuración personalizada:
        >>> config = {
        ...     'fuente_titulo': 'Arial',
        ...     'color_titulo': RGBColor(0xFF, 0x00, 0x00),
        ...     'margin_top': Inches(0.5)
        ... }
        >>> builder = DocumentBuilder(config=config)
        >>> builder.indice() \\
        ...        .titulo("Capítulo 1", 1) \\
        ...        .guardar("documento.docx")
    """

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Inicializa el constructor con configuración personalizable.

        Args:
            config: Diccionario con configuración opcional. Las claves soportadas incluyen:
                - page_height, page_width: Dimensiones de página (Inches)
                - margin_top, margin_bottom, margin_left, margin_right: Márgenes (Inches)
                - fuente_titulo, fuente_texto: Nombres de fuentes
                - color_titulo, color_subtitulo: Colores RGB
                - size_titulo_1, size_titulo_2, size_titulo_3, size_texto: Tamaños (Pt)
                - ancho_tabla_default: Ancho por defecto de tablas (float)
        """
        self.doc = Document()
        self.config = self._configuracion_default()

        if config:
            self.config.update(config)

        self._configurar_pagina()
        self._historial: List[str] = []

    def _configuracion_default(self) -> Dict[str, Any]:
        """
        Retorna la configuración por defecto del documento.

        Returns:
            Diccionario con todos los parámetros de configuración por defecto.
        """
        return {
            # Dimensiones página A4
            'page_height': Inches(11.69),
            'page_width': Inches(8.27),
            'margin_top': Inches(1),
            'margin_bottom': Inches(1),
            'margin_left': Inches(1),
            'margin_right': Inches(1),

            # Fuentes
            'fuente_titulo': 'Lora',
            'fuente_texto': 'Segoe UI Light',

            # Colores corporativos
            'color_titulo': RGBColor(0x2E, 0x3F, 0x5F),
            'color_subtitulo': RGBColor(0x4F, 0x4F, 0x4F),

            # Tamaños de fuente
            'size_titulo_1': Pt(14),
            'size_titulo_2': Pt(12),
            'size_titulo_3': Pt(11),
            'size_texto': Pt(8),
            'size_pie_figura': Pt(6),

            # Tablas
            'ancho_tabla_default': 6.0,
            'size_tabla_header': Pt(6.5),
            'size_tabla_datos': Pt(7),
        }

    def _configurar_pagina(self) -> None:
        """Aplica la configuración de página al documento."""
        section = self.doc.sections[0]
        section.page_height = self.config['page_height']
        section.page_width = self.config['page_width']
        section.top_margin = self.config['margin_top']
        section.bottom_margin = self.config['margin_bottom']
        section.left_margin = self.config['margin_left']
        section.right_margin = self.config['margin_right']

    def titulo(self, texto: str, nivel: int = 1) -> 'DocumentBuilder':
        """
        Agrega un título al documento con formato jerárquico.

        Args:
            texto: Texto del título
            nivel: Nivel jerárquico (1=principal, 2=secundario, 3=terciario)

        Returns:
            self para permitir encadenamiento de métodos

        Example:
            >>> builder.titulo("Capítulo 1", 1)
            >>> builder.titulo("Sección 1.1", 2)
            >>> builder.titulo("Subsección 1.1.1", 3)
        """
        agregar_titulo(self.doc, texto, nivel)
        self._historial.append(f"Título nivel {nivel}: {texto}")
        return self

    def parrafo(self, texto: str) -> 'DocumentBuilder':
        """
        Agrega un párrafo justificado al documento.

        Args:
            texto: Contenido del párrafo

        Returns:
            self para permitir encadenamiento de métodos
        """
        agregar_parrafo(self.doc, texto)
        self._historial.append(f"Párrafo: {texto[:50]}..." if len(texto) > 50 else f"Párrafo: {texto}")
        return self

    def tabla(self, df, titulo: Optional[str] = None, con_merge: bool = False,
              group_cols: Optional[List[str]] = None) -> 'DocumentBuilder':
        """
        Inserta una tabla desde un DataFrame de pandas.

        Args:
            df: DataFrame con los datos
            titulo: Título opcional para la tabla
            con_merge: Si True, usa insertar_tabla_con_merge
            group_cols: Columnas para agrupar (solo si con_merge=True)

        Returns:
            self para permitir encadenamiento de métodos

        Raises:
            ValueError: Si df está vacío o no es válido
        """
        try:
            if df.empty:
                self.parrafo("[Tabla vacía: sin datos para mostrar]")
                self._historial.append("Tabla: vacía")
                return self

            if con_merge:
                insertar_tabla_con_merge(self.doc, df, titulo, group_cols)
                self._historial.append(f"Tabla con merge: {len(df)} filas × {len(df.columns)} cols")
            else:
                insertar_tabla(self.doc, df, titulo)
                self._historial.append(f"Tabla: {len(df)} filas × {len(df.columns)} cols")
        except Exception as e:
            self.parrafo(f"[Error al insertar tabla: {str(e)}]")
            self._historial.append(f"Error en tabla: {str(e)}")

        return self

    def figura(self, figura, titulo: Optional[str] = None,
               pie: Optional[str] = None, ancho_cm: Optional[float] = None, alto_cm: Optional[float] = None) -> 'DocumentBuilder':
        """
        Inserta una figura (matplotlib) en el documento.

        Args:
            figura: Objeto Figure de matplotlib
            titulo: Título opcional sobre la figura
            pie: Texto de pie de figura

        Returns:
            self para permitir encadenamiento de métodos
        """
        try:
            insertar_figura(self.doc, figura, titulo, pie, ancho_cm, alto_cm)
            self._historial.append(f"Figura: {titulo if titulo else 'sin título'}")
        except Exception as e:
            self.parrafo(f"[Error al insertar figura: {str(e)}]")
            self._historial.append(f"Error en figura: {str(e)}")

        return self

    def vinetas(self, items: List[str], nivel: int = 1,
                espacio_antes: Pt = Pt(4),
                espacio_despues: Pt = Pt(4)) -> 'DocumentBuilder':
        """
        Agrega una lista con viñetas (guiones).

        Args:
            items: Lista de textos para las viñetas
            nivel: Nivel de indentación (1=sin indent, 2=indent, etc.)
            espacio_antes: Espaciado antes de cada ítem
            espacio_despues: Espaciado después de cada ítem

        Returns:
            self para permitir encadenamiento de métodos
        """
        agregar_viñetas(self.doc, items, nivel, espacio_antes, espacio_despues)
        self._historial.append(f"Viñetas: {len(items)} items (nivel {nivel})")
        return self

    def salto_pagina(self) -> 'DocumentBuilder':
        """
        Inserta un salto de página.

        Returns:
            self para permitir encadenamiento de métodos
        """
        insertar_salto_pagina(self.doc)
        self._historial.append("Salto de página")
        return self

    def indice(self, titulo: str = "Índice") -> 'DocumentBuilder':
        """
        Inserta una tabla de contenidos (TOC) automática.

        Args:
            titulo: Título de la sección de índice

        Returns:
            self para permitir encadenamiento de métodos

        Note:
            El índice debe actualizarse manualmente en Word (clic derecho > Actualizar campos)
        """
        insertar_indice(self.doc, titulo)
        self._historial.append(f"Índice: {titulo}")
        return self

    def advertencia_actualizacion(self) -> 'DocumentBuilder':
        """
        Agrega advertencia sobre actualización de campos en Word.

        Returns:
            self para permitir encadenamiento de métodos
        """
        agregar_advertencia_actualizacion(self.doc)
        self._historial.append("Advertencia de actualización")
        return self

    def numerar_titulos(self) -> 'DocumentBuilder':
        """
        Numera automáticamente todos los títulos existentes (1., 1.1, 1.1.1, etc.).

        Returns:
            self para permitir encadenamiento de métodos
        """
        numerar_titulos_existentes(self.doc)
        self._historial.append("Numeración de títulos aplicada")
        return self

    def guardar(self, ruta: str, verbose: bool = True) -> None:
        """
        Guarda el documento en la ruta especificada.

        Args:
            ruta: Ruta del archivo .docx de salida
            verbose: Si True, imprime información del guardado
        """
        try:
            guardar_documento(self.doc, ruta)
            if verbose:
                print(f"[OK] Documento guardado en: {ruta}")
                print(f"[OK] Operaciones realizadas: {len(self._historial)}")
        except Exception as e:
            print(f"[ERROR] Error al guardar documento: {e}")
            raise

    def obtener_historial(self) -> List[str]:
        """
        Retorna una copia del historial de operaciones.

        Returns:
            Lista de strings describiendo cada operación realizada
        """
        return self._historial.copy()

    def mostrar_historial(self) -> 'DocumentBuilder':
        """
        Imprime el historial de operaciones en consola.

        Returns:
            self para permitir encadenamiento de métodos
        """
        print("\n=== Historial de Operaciones ===")
        for i, op in enumerate(self._historial, 1):
            print(f"{i:2d}. {op}")
        print(f"\nTotal: {len(self._historial)} operaciones\n")
        return self

    def obtener_documento(self) -> Document:
        """
        Retorna el objeto Document interno para manipulación avanzada.

        Returns:
            Objeto Document de python-docx

        Warning:
            Modificar directamente el documento puede afectar el historial
        """
        return self.doc


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Generador de documento Word (A4) con utilidades de formato.")
    parser.add_argument("salida", help="Ruta del archivo .docx de salida, p.ej.: informe.docx")
    parser.add_argument("--builder", action="store_true", help="Usar DocumentBuilder en lugar de funciones")
    args = parser.parse_args()

    if args.builder:
        # EJEMPLO CON DOCUMENTBUILDER (Método recomendado)
        print("Generando documento con DocumentBuilder...")

        builder = DocumentBuilder()
        builder.indice() \
               .advertencia_actualizacion() \
               .titulo("Resumen de Proyecto", 1) \
               .parrafo("Este es un ejemplo mínimo de documento generado con DocumentBuilder.") \
               .titulo("Sección Principal", 2) \
               .vinetas(["Ítem 1", "Ítem 2", "Ítem 3"]) \
               .salto_pagina() \
               .titulo("Conclusiones", 1) \
               .parrafo("El DocumentBuilder permite crear documentos de forma más limpia y fluida.") \
               .mostrar_historial() \
               .guardar(args.salida)
    else:
        # EJEMPLO CON FUNCIONES (Método tradicional)
        print("Generando documento con funciones tradicionales...")

        doc = crear_documento_a4()
        insertar_indice(doc)
        agregar_advertencia_actualizacion(doc)
        agregar_titulo(doc, "Resumen de Proyecto", 1)
        agregar_parrafo(doc, "Este es un ejemplo mínimo de documento generado desde funciones.")
        agregar_titulo(doc, "Sección", 2)
        agregar_viñetas(doc, ["Ítem 1", "Ítem 2", "Ítem 3"])

        guardar_documento(doc, args.salida)
        print(f"Documento guardado en: {args.salida}")


